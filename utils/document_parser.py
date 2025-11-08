# import logging
# import os
# from pathlib import Path
# from typing import List, Dict

# logger = logging.getLogger(__name__)

# SUPPORTED_FORMATS = {".pdf", ".docx", ".md", ".txt"}

# def parse_document(file_path: str) -> str:
#     logger.info(f"解析支持的文档格式并返回纯文本:{'、'.join(SUPPORTED_FORMATS)}")
#     if not os.path.exists(file_path):
#         logger.error(f"文件不存在: {file_path}")
#         raise FileNotFoundError(f"File not found: {file_path}")

#     ext = Path(file_path).suffix.lower()
#     if ext not in SUPPORTED_FORMATS:
#         logger.error(f"不支持的文件格式: {ext}")
#         raise ValueError(f"Unsupported file format: {ext}")

#     try:
#         if ext == ".pdf":
#             return _parse_pdf(file_path)
#         elif ext == ".docx":
#             return _parse_docx(file_path)
#         elif ext == ".md":
#             return _parse_markdown(file_path)
#         elif ext == ".txt":
#             return _parse_txt(file_path)
#     except Exception as e:
#         logger.error(f"解析文件失败: {file_path}, 错误: {e}")
#         raise

# def _parse_pdf(file_path: str) -> str:
#     from PyPDF2 import PdfReader
#     reader = PdfReader(file_path)
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text()
#     logger.info(f"成功解析 PDF 文件: {file_path}")
#     return text

# def _parse_docx(file_path: str) -> str:
#     from docx import Document
#     doc = Document(file_path)
#     text = "\n".join([para.text for para in doc.paragraphs])
#     logger.info(f"成功解析 Word 文件: {file_path}")
#     return text

# def _parse_markdown(file_path: str) -> str:
#     with open(file_path, 'r', encoding='utf-8') as f:
#         text = f.read()
#     logger.info(f"成功解析 Markdown 文件: {file_path}")
#     return text

# def _parse_txt(file_path: str) -> str:
#     with open(file_path, 'r', encoding='utf-8') as f:
#         text = f.read()
#     logger.info(f"成功解析 TXT 文件: {file_path}")
#     return text

# def split_text_into_chunks(text: str, chunk_size: int = 512) -> List[str]:
#     """按字符数切分文本为块"""
#     chunks = []
#     for i in range(0, len(text), chunk_size):
#         chunks.append(text[i:i + chunk_size])
#     logger.debug(f"文本切分为 {len(chunks)} 个块")
#     return chunks


import logging
import os
from pathlib import Path
from typing import List, Dict, Any
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument
import openai as OpenAI
from dotenv import load_dotenv

logger = logging.getLogger(__name__)
load_dotenv()

SUPPORTED_FORMATS = {".pdf", ".docx", ".md", ".txt"}

def parse_document(file_path: str) -> str:
    logging.info('解析文档为纯文本')
    logger.info(f"解析支持的文档格式: {'、'.join(SUPPORTED_FORMATS)}")
    if not os.path.exists(file_path):
        logger.error(f"文件不存在: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = Path(file_path).suffix.lower()
    if ext not in SUPPORTED_FORMATS:
        logger.error(f"不支持的文件格式: {ext}")
        raise ValueError(f"Unsupported file format: {ext}")

    try:
        if ext == ".pdf":
            logger.info(f"正在解析 PDF 文件: {file_path}")
            return _parse_pdf(file_path)
        elif ext == ".docx":
            logger.info(f"正在解析 DOCX 文件: {file_path}")
            return _parse_docx(file_path)
        elif ext == ".md":
            logger.info(f"正在解析 MD 文件: {file_path}")
            return _parse_markdown(file_path)
        elif ext == ".txt":
            logger.info(f"正在解析 TXT 文件: {file_path}")
            return _parse_txt(file_path)
    except Exception as e:
        logger.error(f"解析文件失败: {file_path}, 错误: {e}")
        raise

# === 以下解析函数保持不变 ===
def _parse_pdf(file_path: str) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(file_path)
    text = ""
    for page_num, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n\n"
    logger.info(f"成功解析 PDF 文件: {file_path} ({len(reader.pages)} 页)")
    return text

def _parse_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    logger.info(f"成功解析 Word 文件: {file_path}")
    return text

def _parse_markdown(file_path: str) -> str:
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    logger.info(f"成功解析 Markdown 文件: {file_path}")
    return text

def _parse_txt(file_path: str) -> str:
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    logger.info(f"成功解析 TXT 文件: {file_path}")
    return text


# ==============================
# 新增：语义分段 + 向量化
# ==============================

def create_documents_with_metadata(file_path: str, chunk_size: int = 512, chunk_overlap: int = 80) -> List[LCDocument]:
    """
    解析文档并切分为带元数据的 LangChain Document 列表
    """
    raw_text = parse_document(file_path)
    file_name = os.path.basename(file_path)

    # 使用递归分段器（优先按 \n\n → \n → " " 切分，避免在句子中间切断）
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", "；", " ", ""],
        length_function=len,  # 按字符数（中文适用）
        keep_separator=True
    )

    chunks = text_splitter.split_text(raw_text)
    documents = []
    for i, chunk in enumerate(chunks):
        if chunk.strip():  # 跳过空块
            doc = LCDocument(
                page_content=chunk.strip(),
                metadata={
                    "source": file_path,
                    "file_name": file_name,
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
            )
            documents.append(doc)

    logger.info(f"文档 {file_name} 切分为 {len(documents)} 个语义块")
    return documents


def embed_documents(documents: List[LCDocument]) -> List[Dict[str, Any]]:
    """
    对文档列表进行向量化（使用 DashScope Qwen Embedding）
    返回: [{"content": "...", "embedding": [...], "metadata": {...}}, ...]
    """
    api_key = os.getenv("qianwen_api_key")
    if not api_key:
        raise ValueError("请设置环境变量 qianwen_api_key")

    client = OpenAI(
        api_key=api_key,
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
    )

    # 提取所有文本
    texts = [doc.page_content for doc in documents]
    
    try:
        # 调用 embedding 接口（DashScope 兼容 OpenAI /embeddings）
        response = client.embeddings.create(
            model="text-embedding-v2",  # 阿里云 Qwen Embedding 模型
            input=texts
        )
        
        results = []
        for i, doc in enumerate(documents):
            emb = response.data[i].embedding
            results.append({
                "content": doc.page_content,
                "embedding": emb,
                "metadata": doc.metadata
            })
        
        logger.info(f"成功向量化 {len(results)} 个文档块")
        return results

    except Exception as e:
        logger.error(f"向量化失败: {e}")
        raise